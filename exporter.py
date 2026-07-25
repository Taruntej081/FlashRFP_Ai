import re
import io
import difflib
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx_stream(question, response_text, sources=None):
    """
    Generates a professionally styled Microsoft Word document (.docx) as a byte stream.
    Parses basic Markdown formatting (headers, bolding, list items) to native Word styles.
    """
    doc = docx.Document()
    
    # Configure page margins (Standard 1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(1)
        section.bottom_margin = docx.shared.Inches(1)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)

    # 1. Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("FlashRFP AI - Proposal Response")
    title_run.bold = True
    title_run.font.name = "Inter"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(9, 9, 11)  # Charcoal / Zinc 950
    title_p.paragraph_format.space_after = Pt(2)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Automated RFP Bid Response Engine | Draft Document")
    sub_run.font.name = "Inter"
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(113, 113, 122)  # Zinc 500
    sub_run.italic = True
    sub_p.paragraph_format.space_after = Pt(24)

    # 2. RFP Question Section
    q_hdr_p = doc.add_paragraph()
    q_hdr_run = q_hdr_p.add_run("RFP Question:")
    q_hdr_run.bold = True
    q_hdr_run.font.name = "Inter"
    q_hdr_run.font.size = Pt(11)
    q_hdr_run.font.color.rgb = RGBColor(37, 99, 235)  # Accent Blue
    q_hdr_p.paragraph_format.space_after = Pt(4)

    q_body_p = doc.add_paragraph()
    q_body_run = q_body_p.add_run(question)
    q_body_run.bold = True
    q_body_run.font.name = "Inter"
    q_body_run.font.size = Pt(13)
    q_body_run.font.color.rgb = RGBColor(9, 9, 11)
    q_body_p.paragraph_format.space_after = Pt(18)

    # Horizontal Divider Line
    divider_p = doc.add_paragraph()
    divider_p.paragraph_format.space_after = Pt(18)
    # Add a border-like run using underscores
    div_run = divider_p.add_run("_" * 50)
    div_run.font.color.rgb = RGBColor(228, 228, 231)  # Zinc 200
    div_run.font.size = Pt(10)

    # 3. Drafted Response Section
    a_hdr_p = doc.add_paragraph()
    a_hdr_run = a_hdr_p.add_run("Drafted Bid Response:")
    a_hdr_run.bold = True
    a_hdr_run.font.name = "Inter"
    a_hdr_run.font.size = Pt(11)
    a_hdr_run.font.color.rgb = RGBColor(37, 99, 235)  # Accent Blue
    a_hdr_p.paragraph_format.space_after = Pt(8)

    # Parse response text into styled paragraphs
    lines = response_text.split("\n")
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue
            
        # Detect bold wrappers and strip them for detection
        clean_detect = stripped_line
        if clean_detect.startswith("**") and clean_detect.endswith("**"):
            clean_detect = clean_detect[2:-2].strip()

        # Headers (e.g. ## Heading)
        if clean_detect.startswith("#"):
            level = len(clean_detect) - len(clean_detect.lstrip("#"))
            header_text = clean_detect.lstrip("#").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            size = 14 if level == 1 else (12.5 if level == 2 else 11.5)
            _add_inline_formatted_text(p, header_text, size_pt=size, bold_default=True)
            
        # Bullet list items (e.g. * item or - item)
        elif stripped_line.startswith("* ") or stripped_line.startswith("- ") or stripped_line.startswith("• "):
            list_text = stripped_line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            _add_inline_formatted_text(p, list_text)
            
        # Numbered section headers (e.g. 1. Introduction or **1. Introduction**)
        elif clean_detect[0].isdigit() and "." in clean_detect[:4] and len(clean_detect) < 80:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            _add_inline_formatted_text(p, clean_detect, size_pt=12, bold_default=True)

        # Numbered list items (e.g. 1. item)
        elif re.match(r"^\d+\.\s+", clean_detect):
            match = re.match(r"^(\d+)\.\s+(.*)", clean_detect)
            num_text = match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            _add_inline_formatted_text(p, num_text)
            
        # Standard paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.15
            _add_inline_formatted_text(p, stripped_line)

    # 4. References / Source Citations Section
    if sources:
        doc.add_page_break()
        
        ref_title_p = doc.add_paragraph()
        ref_title_p.paragraph_format.space_before = Pt(24)
        ref_title_p.paragraph_format.space_after = Pt(12)
        ref_title_run = ref_title_p.add_run("References & Source Citations")
        ref_title_run.bold = True
        ref_title_run.font.name = "Inter"
        ref_title_run.font.size = Pt(14)
        ref_title_run.font.color.rgb = RGBColor(113, 113, 122)  # Zinc 500
        
        info_p = doc.add_paragraph()
        info_run = info_p.add_run(
            "The following snippets from your historical proposal documents were used "
            "as the source of truth to synthesize the drafted response above:"
        )
        info_run.font.name = "Inter"
        info_run.font.size = Pt(9.5)
        info_run.font.color.rgb = RGBColor(113, 113, 122)
        info_p.paragraph_format.space_after = Pt(16)
        
        for idx, src in enumerate(sources):
            src_p = doc.add_paragraph()
            src_p.paragraph_format.space_after = Pt(12)
            
            # Citation Header
            run_idx = src_p.add_run(f"[{idx+1}] File: ")
            run_idx.bold = True
            run_idx.font.name = "Inter"
            run_idx.font.size = Pt(10)
            run_idx.font.color.rgb = RGBColor(37, 99, 235)
            
            run_file = src_p.add_run(f"{src['source']}  |  Similarity Score: {src['similarity']}\n")
            run_file.bold = True
            run_file.font.name = "Inter"
            run_file.font.size = Pt(10)
            run_file.font.color.rgb = RGBColor(9, 9, 11)
            
            # Citation Snippet
            run_snippet = src_p.add_run(f"\"{src['text'].strip()}\"")
            run_snippet.italic = True
            run_snippet.font.name = "Inter"
            run_snippet.font.size = Pt(9)
            run_snippet.font.color.rgb = RGBColor(82, 82, 91)  # Zinc 600

    # 5. Output Byte Stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

def _add_inline_formatted_text(paragraph, text, size_pt=10.5, font_name="Inter", bold_default=False):
    """
    Helper to parse basic inline markdown formatting like **bold** and add it to the paragraph.
    Removes the literal '**' characters from the rendered Word document.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            clean_part = part[2:-2]
            run = paragraph.add_run(clean_part)
            run.bold = True
        else:
            run = paragraph.add_run(part)
            if bold_default:
                run.bold = True
            
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(9, 9, 11)

def find_best_matching_answer_with_score(doc_question, qa_pairs):
    """
    Given a question text, find the best matching question in qa_pairs
    and return the answer and similarity score.
    """
    if not doc_question or not qa_pairs:
        return None, 0.0
        
    doc_question_clean = doc_question.strip().lower()
    
    # Clean common prepended headers/numbers (e.g. "1.1 ", "Section A - ")
    cleaned_candidate = re.sub(r'^(section\s+[a-z0-9.-]+|question\s+[0-9.]+|[0-9.-]+)\s*[:.-]?\s*', '', doc_question_clean).strip()
    
    # Try exact or clean substring matching first
    for pair in qa_pairs:
        q_clean = pair["question"].strip().lower()
        q_clean_noprefix = re.sub(r'^(section\s+[a-z0-9.-]+|question\s+[0-9.]+|[0-9.-]+)\s*[:.-]?\s*', '', q_clean).strip()
        if (cleaned_candidate and q_clean_noprefix and 
            (cleaned_candidate in q_clean_noprefix or q_clean_noprefix in cleaned_candidate)):
            return pair["answer"], 1.0
            
    # Fallback to fuzzy matching
    best_match = None
    best_score = 0.0
    for pair in qa_pairs:
        q_clean = pair["question"].strip().lower()
        score = difflib.SequenceMatcher(None, doc_question_clean, q_clean).ratio()
        if score > best_score:
            best_score = score
            best_match = pair
            
    if best_match:
        return best_match["answer"], best_score
        
    return None, 0.0

def find_best_matching_answer(doc_question, qa_pairs, threshold=0.5):
    """
    Given a question text from a document, find the best matching question in qa_pairs
    and return its answer.
    """
    ans, score = find_best_matching_answer_with_score(doc_question, qa_pairs)
    if score >= threshold:
        return ans
    return None

def fill_rfp_docx_template(template_path_or_stream, qa_pairs, threshold=0.5):
    """
    Fills in the answers into a blank Word RFP document.
    Replaces paragraph placeholders, text box placeholders, and empty/placeholder table cells
    while locking original Word document formatting, fonts, and styles.
    """
    doc = docx.Document(template_path_or_stream)
    
    recent_paragraphs = []
    
    def is_placeholder_text(text):
        t_lower = text.lower().strip()
        if not t_lower:
            return False
        placeholders = [
            "insert answer", "insert response", "answer here", "response here", 
            "insert response here", "insert answer here", "[answer]", "[response]",
            "vendor response", "vendor answer", "bidder response", "supplier response",
            "your response", "tbd", "to be filled", "write answer here", "specify details",
            "fill here", "compliance status", "[insert]", "[write here]"
        ]
        if any(p in t_lower for p in placeholders):
            return True
        if re.search(r'\[[^\]]*(insert|answer|response|vendor|bidder|write|here|tbd)[^\]]*\]', t_lower, re.IGNORECASE):
            return True
        if re.search(r'\{[^}]*(insert|answer|response|vendor|bidder|write|here|tbd)[^}]*\}', t_lower, re.IGNORECASE):
            return True
        if re.search(r'\<[^>]*(insert|answer|response|vendor|bidder|write|here|tbd)[^>]*\>', t_lower, re.IGNORECASE):
            return True
        return False
        
    def replace_placeholder_in_paragraph(paragraph, target_text, replacement_text):
        full_text = "".join(run.text for run in paragraph.runs)
        font_name = "Inter"
        size_pt = 10.5
        if paragraph.runs:
            if paragraph.runs[0].font.name:
                font_name = paragraph.runs[0].font.name
            if paragraph.runs[0].font.size:
                size_pt = paragraph.runs[0].font.size.pt
                
        if target_text in full_text:
            paragraph.text = ""
            _add_inline_formatted_text(paragraph, replacement_text, size_pt=size_pt, font_name=font_name)
        else:
            paragraph.text = ""
            _add_inline_formatted_text(paragraph, replacement_text, size_pt=size_pt, font_name=font_name)

    # 1. Fill Paragraph Placeholders
    for para in doc.paragraphs:
        p_text = para.text.strip()
        if not p_text:
            continue
            
        if is_placeholder_text(p_text):
            target = p_text
            match = re.search(r'\[[^\]]*(insert|answer|response|write|here|tbd)[^\]]*\]', p_text, re.IGNORECASE)
            if match:
                target = match.group(0)
                
            # Look back to find the question
            question_candidate = ""
            for prev_text in reversed(recent_paragraphs):
                if prev_text and not is_placeholder_text(prev_text):
                    question_candidate = prev_text
                    break
            
            if question_candidate:
                answer = find_best_matching_answer(question_candidate, qa_pairs, threshold)
                if answer:
                    replace_placeholder_in_paragraph(para, target, answer)
        else:
            recent_paragraphs.append(p_text)
            if len(recent_paragraphs) > 10:
                recent_paragraphs.pop(0)
                
    # 2. Fill Tables (empty/placeholder cells next to or below questions)
    for table in doc.tables:
        for r_idx in range(len(table.rows)):
            row = table.rows[r_idx]
            for c_idx in range(len(row.cells)):
                cell = row.cells[c_idx]
                cell_text = cell.text.strip()
                
                is_empty = not cell_text
                is_placeholder = is_placeholder_text(cell_text) if not is_empty else False
                
                if is_empty or is_placeholder:
                    # Check cells above and all preceding cells in row
                    best_answer = None
                    highest_score = 0.0

                    # Check cell above
                    if r_idx > 0:
                        cell_above = table.rows[r_idx - 1].cells[c_idx]
                        cand_above = cell_above.text.strip()
                        if cand_above and not is_placeholder_text(cand_above):
                            ans_a, s_a = find_best_matching_answer_with_score(cand_above, qa_pairs)
                            if ans_a and s_a >= threshold and s_a > highest_score:
                                highest_score = s_a
                                best_answer = ans_a

                    # Check preceding cells in row
                    for prev_c in range(c_idx - 1, -1, -1):
                        cell_left = row.cells[prev_c]
                        cand_left = cell_left.text.strip()
                        if cand_left and not is_placeholder_text(cand_left):
                            ans_l, s_l = find_best_matching_answer_with_score(cand_left, qa_pairs)
                            if ans_l and s_l >= threshold and s_l > highest_score:
                                highest_score = s_l
                                best_answer = ans_l
                                break

                    if best_answer:
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _add_inline_formatted_text(p, best_answer, size_pt=10, font_name="Inter")

    # 3. Fill Textbox Placeholders
    try:
        for p_elem in doc.element.xpath('//w:txbxContent//w:p'):
            p_obj = docx.text.paragraph.Paragraph(p_elem, doc)
            p_txt = p_obj.text.strip()
            if is_placeholder_text(p_txt):
                # Search previous paragraphs in body for question
                if recent_paragraphs:
                    q_cand = recent_paragraphs[-1]
                    ans = find_best_matching_answer(q_cand, qa_pairs, threshold)
                    if ans:
                        p_obj.text = ""
                        _add_inline_formatted_text(p_obj, ans, size_pt=9.5, font_name="Inter")
    except Exception:
        pass

    # Return document bytes
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream.getvalue()

def generate_batch_docx_stream(qa_results):
    """
    Generates a beautifully formatted DOCX response document for multiple QA pairs.
    """
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(1)
        section.bottom_margin = docx.shared.Inches(1)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)

    # Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("FlashRFP AI - Complete Proposal Response")
    title_run.bold = True
    title_run.font.name = "Inter"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(9, 9, 11)
    title_p.paragraph_format.space_after = Pt(2)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Automated Batch RFP Response Document")
    sub_run.font.name = "Inter"
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(113, 113, 122)
    sub_run.italic = True
    sub_p.paragraph_format.space_after = Pt(24)

    for idx, qa in enumerate(qa_results):
        question = qa["question"]
        response_text = qa["answer"]
        
        # Section Heading
        sec_p = doc.add_paragraph()
        sec_p.paragraph_format.space_before = Pt(18)
        sec_p.paragraph_format.space_after = Pt(4)
        sec_run = sec_p.add_run(f"Question {idx+1}:")
        sec_run.bold = True
        sec_run.font.name = "Inter"
        sec_run.font.size = Pt(11)
        sec_run.font.color.rgb = RGBColor(37, 99, 235)
        
        q_body_p = doc.add_paragraph()
        q_body_run = q_body_p.add_run(question)
        q_body_run.bold = True
        q_body_run.font.name = "Inter"
        q_body_run.font.size = Pt(12)
        q_body_run.font.color.rgb = RGBColor(9, 9, 11)
        q_body_p.paragraph_format.space_after = Pt(12)
        
        # Response Title
        resp_p = doc.add_paragraph()
        resp_run = resp_p.add_run("Drafted Response:")
        resp_run.bold = True
        resp_run.font.name = "Inter"
        resp_run.font.size = Pt(10.5)
        resp_run.font.color.rgb = RGBColor(113, 113, 122)
        resp_p.paragraph_format.space_after = Pt(6)
        
        # Parse and add markdown response lines
        lines = response_text.split("\n")
        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue
                
            clean_detect = stripped_line
            if clean_detect.startswith("**") and clean_detect.endswith("**"):
                clean_detect = clean_detect[2:-2].strip()

            if clean_detect.startswith("#"):
                level = len(clean_detect) - len(clean_detect.lstrip("#"))
                header_text = clean_detect.lstrip("#").strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                size = 13 if level == 1 else (12 if level == 2 else 11)
                _add_inline_formatted_text(p, header_text, size_pt=size, bold_default=True)
                
            elif stripped_line.startswith("* ") or stripped_line.startswith("- ") or stripped_line.startswith("• "):
                list_text = stripped_line[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                _add_inline_formatted_text(p, list_text)
                
            elif clean_detect[0].isdigit() and "." in clean_detect[:4] and len(clean_detect) < 80:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                _add_inline_formatted_text(p, clean_detect, size_pt=11.5, bold_default=True)

            elif re.match(r"^\d+\.\s+", clean_detect):
                match = re.match(r"^(\d+)\.\s+(.*)", clean_detect)
                num_text = match.group(2).strip()
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(3)
                _add_inline_formatted_text(p, num_text)
                
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                _add_inline_formatted_text(p, stripped_line)
                
        # Divider between sections (except the last)
        if idx < len(qa_results) - 1:
            div_p = doc.add_paragraph()
            div_p.paragraph_format.space_before = Pt(12)
            div_p.paragraph_format.space_after = Pt(12)
            div_run = div_p.add_run("_" * 50)
            div_run.font.name = "Inter"
            div_run.font.color.rgb = RGBColor(228, 228, 231)
            div_run.font.size = Pt(10)

    # Output Byte Stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

def process_boq_excel(file_bytes_or_stream, collection, api_key, tenant_id=None, provider=None, model=None, win_themes=None, demo_mode=False, progress_callback=None):
    """
    Parses a government/tender Excel BOQ (Bill of Quantities / Technical Compliance) file.
    Auto-detects item descriptions, queries knowledge base (ChromaDB), calls AI to evaluate compliance
    and generate technical responses, then writes responses directly back into the Excel spreadsheet.
    
    Returns:
        tuple: (output_excel_bytes: bytes, rows_processed_list: list[dict])
    """
    import json
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from rag_engine import query_knowledge_base, call_llm

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes_or_stream))
    ws = wb.active

    # 1. Identify Header Row and Column Indices
    header_row_idx = 1
    desc_col_idx = None
    comp_col_idx = None
    resp_col_idx = None
    remarks_col_idx = None

    desc_keywords = ["requirement", "requirements", "description", "specification", "particulars", "item name", "technical requirement", "equipment", "boq item", "details", "tender spec", "item description"]
    comp_keywords = ["compliance", "complied", "compliance status", "vendor compliance", "status"]
    resp_keywords = ["proposed", "offered", "vendor specification", "ai response", "technical spec", "model offered", "ai technical response", "response"]
    remarks_keywords = ["remarks", "justification", "deviation", "ai remarks", "reference", "notes"]

    max_r = min(ws.max_row, 30)
    max_c = ws.max_column

    # Search for header row
    for r in range(1, max_r + 1):
        row_vals = [str(ws.cell(row=r, column=c).value or "").strip().lower() for c in range(1, max_c + 1)]
        if any(any(kw in val for kw in desc_keywords) for val in row_vals):
            header_row_idx = r
            for c_idx, val in enumerate(row_vals, 1):
                if not desc_col_idx and any(kw in val for kw in desc_keywords):
                    desc_col_idx = c_idx
                elif not comp_col_idx and any(kw in val for kw in comp_keywords):
                    comp_col_idx = c_idx
                elif not resp_col_idx and any(kw in val for kw in resp_keywords):
                    resp_col_idx = c_idx
                elif not remarks_col_idx and any(kw in val for kw in remarks_keywords):
                    remarks_col_idx = c_idx
            break

    # Default description column to col 2 if not auto-detected
    if not desc_col_idx:
        desc_col_idx = 2 if max_c >= 2 else 1

    # Dynamically append columns if not present
    new_col_cursor = max_c
    header_fill = PatternFill(start_color="10B981", end_color="10B981", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")

    if not comp_col_idx:
        new_col_cursor += 1
        comp_col_idx = new_col_cursor
        cell = ws.cell(row=header_row_idx, column=comp_col_idx)
        cell.value = "AI Compliance Status"
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    if not resp_col_idx:
        new_col_cursor += 1
        resp_col_idx = new_col_cursor
        cell = ws.cell(row=header_row_idx, column=resp_col_idx)
        cell.value = "AI Technical Specification & Response"
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    if not remarks_col_idx:
        new_col_cursor += 1
        remarks_col_idx = new_col_cursor
        cell = ws.cell(row=header_row_idx, column=remarks_col_idx)
        cell.value = "AI Remarks & Sourced Context"
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # 2. Iterate Rows and Process via RAG + LLM
    rows_processed = []
    total_data_rows = max(1, ws.max_row - header_row_idx)
    processed_count = 0

    # Styling fills for compliance
    fill_complied = PatternFill(start_color="D1FAE5", end_color="D1FAE5", fill_type="solid")
    font_complied = Font(name="Calibri", size=10, bold=True, color="065F46")

    fill_dev = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    font_dev = Font(name="Calibri", size=10, bold=True, color="92400E")

    fill_noncomp = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")
    font_noncomp = Font(name="Calibri", size=10, bold=True, color="991B1B")

    body_font = Font(name="Calibri", size=10)
    remarks_font = Font(name="Calibri", size=9, italic=True)

    for r_idx in range(header_row_idx + 1, ws.max_row + 1):
        item_cell = ws.cell(row=r_idx, column=desc_col_idx)
        item_text = str(item_cell.value or "").strip()

        if not item_text or item_text.lower() in ["total", "subtotal", "grand total", "tax", "gst", "sl no", "s.no", "item"]:
            continue

        processed_count += 1
        if progress_callback:
            progress_callback(processed_count, total_data_rows, item_text)

        # Query Knowledge Base
        contexts = query_knowledge_base(item_text, collection, top_k=5, tenant_id=tenant_id)
        context_str = "\n".join([f"- {c['source']} (Score: {c['similarity']}): {c['text']}" for c in contexts]) if contexts else "No historical proposal matches found."

        compliance_val = "COMPLIED"
        tech_resp = ""
        remarks_val = ""

        if demo_mode:
            compliance_val = "COMPLIED"
            tech_resp = f"Offered specification compliant with requirement: '{item_text}'. Equipped with enterprise high-availability hardware standards."
            remarks_val = f"Matched historical KB context ({len(contexts)} sources)."
        else:
            prompt = f"""
            You are an expert technical bid evaluator and proposal engineer.
            Analyze the following BOQ Item / Specification and evaluate compliance based on historical knowledge base context.

            BOQ ITEM / SPECIFICATION:
            "{item_text}"

            HISTORICAL KNOWLEDGE BASE CONTEXT:
            {context_str}

            WIN THEMES / VALUE PROPOSITION:
            {win_themes or "High performance, enterprise reliability, 24/7 SLA"}

            Return JSON object with EXACT keys:
            1. "compliance": Must be strictly one of: "COMPLIED", "NON-COMPLIED", or "COMPLIED WITH DEVIATIONS".
            2. "technical_response": Concise, high-precision technical specification matching the requirement.
            3. "remarks": Short 1-sentence justification referencing retrieved context.
            """
            try:
                raw_out = call_llm(api_key, prompt, provider=provider, model=model, json_mode=True)
                json_res = json.loads(raw_out)
                compliance_val = str(json_res.get("compliance", "COMPLIED")).upper().strip()
                tech_resp = str(json_res.get("technical_response", "")).strip()
                remarks_val = str(json_res.get("remarks", "")).strip()
            except Exception as ex:
                compliance_val = "COMPLIED"
                tech_resp = f"Specifications proposed in alignment with: {item_text}."
                remarks_val = f"Generated based on retrieved context. Error: {str(ex)}"

        # Write to openpyxl cells
        c_cell = ws.cell(row=r_idx, column=comp_col_idx)
        c_cell.value = compliance_val
        if "NON" in compliance_val:
            c_cell.fill = fill_noncomp
            c_cell.font = font_noncomp
        elif "DEV" in compliance_val:
            c_cell.fill = fill_dev
            c_cell.font = font_dev
        else:
            c_cell.fill = fill_complied
            c_cell.font = font_complied
        c_cell.alignment = Alignment(horizontal="center", vertical="center")

        r_cell = ws.cell(row=r_idx, column=resp_col_idx)
        r_cell.value = tech_resp
        r_cell.font = body_font
        r_cell.alignment = Alignment(wrap_text=True, vertical="center")

        m_cell = ws.cell(row=r_idx, column=remarks_col_idx)
        m_cell.value = remarks_val
        m_cell.font = remarks_font
        m_cell.alignment = Alignment(wrap_text=True, vertical="center")

        rows_processed.append({
            "row_num": r_idx,
            "item": item_text,
            "compliance": compliance_val,
            "response": tech_resp,
            "remarks": remarks_val
        })

    # Adjust Column Widths for readability
    ws.column_dimensions[openpyxl.utils.get_column_letter(desc_col_idx)].width = 40
    ws.column_dimensions[openpyxl.utils.get_column_letter(comp_col_idx)].width = 25
    ws.column_dimensions[openpyxl.utils.get_column_letter(resp_col_idx)].width = 50
    ws.column_dimensions[openpyxl.utils.get_column_letter(remarks_col_idx)].width = 35

    out_stream = io.BytesIO()
    wb.save(out_stream)
    out_stream.seek(0)
    return out_stream.getvalue(), rows_processed

def generate_advanced_proposal_docx(extracted_data, project_name="RFP Response"):
    """
    Generates a perfectly formatted, enterprise-grade Word document.
    extracted_data: A list of dictionaries containing the AI's output.
    Example: [{"question": "1. Describe SLA...", "answer": "We comply...", "source": "Past_Proposal.pdf", "manual_review": False}]
    """
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 1. Set Default Document Styles (Enterprise Standard)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # 2. Document Title Page
    title = doc.add_heading(project_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Technical & Compliance Proposal Draft")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80) # Gray
    sub_run.italic = True
    
    doc.add_page_break()

    # 3. Process Each Extracted Question & Answer
    for item in extracted_data:
        question = item.get('question', 'Unknown Question')
        answer = item.get('answer', 'No answer generated.')
        source = item.get('source', 'No source cited.')
        is_manual_review = item.get('manual_review', False)

        # Add the Question as a Heading (Level 2)
        doc.add_heading(question, level=2)

        # Add the AI Answer
        answer_para = doc.add_paragraph()
        answer_run = answer_para.add_run(answer)
        
        # SECURITY/REVIEW FEATURE: If the AI flagged it for manual review, make it RED
        if is_manual_review or "MANUAL REVIEW REQUIRED" in answer.upper():
            answer_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Red
            answer_run.bold = True
        else:
            answer_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00) # Black

        # Add the Source Citation (Gray, small, italic)
        source_para = doc.add_paragraph()
        source_run = source_para.add_run(f"Source Reference: {source}")
        source_run.font.size = Pt(9)
        source_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80) # Gray
        source_run.italic = True
        
        # Add spacing between questions
        doc.add_paragraph("")

    # 4. Save to memory buffer for Streamlit download
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def render_export_button(extracted_data):
    """
    Renders Streamlit UI button for advanced enterprise Word proposal export with red manual review flags.
    """
    import streamlit as st
    st.markdown("---")
    st.subheader("📄 Export Proposal Draft")
    
    if st.button("⚡ Generate Formatted Word Document", type="primary", use_container_width=True, key="gen_adv_docx_btn"):
        if extracted_data:
            with st.spinner("Formatting your enterprise Word document..."):
                # Generate the file
                word_file = generate_advanced_proposal_docx(extracted_data)
                
                # Provide the download button
                st.success("✅ Document generated successfully!")
                st.download_button(
                    label="⬇️ Download Proposal Draft (.docx)",
                    data=word_file,
                    file_name="FlashRFP_Proposal_Draft.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_adv_docx_btn"
                )
        else:
            st.warning("Please extract the RFP questions first.")

