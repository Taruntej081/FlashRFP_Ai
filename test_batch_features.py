import os
import sys
import docx
import io
from dotenv import load_dotenv
from rag_engine import (
    chunk_text,
    get_chroma_client,
    get_or_create_collection,
    ingest_documents_batch,
    batch_process_rfp_questions,
    extract_questions_from_pdf
)
from exporter import fill_rfp_docx_template, find_best_matching_answer, generate_batch_docx_stream

def create_mock_docx_template(path):
    """
    Creates a mock Word document template containing:
    1. A question paragraph followed by a paragraph containing a placeholder [Insert Answer Here].
    2. A table with vertical layout (questions in row 0, empty/placeholder cells in row 1).
    3. A table with horizontal layout (questions in col 0, empty cell in col 1).
    """
    doc = docx.Document()
    
    # 1. Paragraph Placeholder
    doc.add_heading("Section 1: Information Security", level=1)
    doc.add_paragraph("What is your security protocol?")
    doc.add_paragraph("[Insert Answer Here]")
    
    doc.add_paragraph("What is your uptime SLA?")
    doc.add_paragraph("[Response Here]")
    
    # 2. Table Vertical Layout (Question cell above)
    doc.add_heading("Section 2: SLA & Availability Table", level=2)
    t1 = doc.add_table(rows=2, cols=2)
    # Row 0: Questions
    t1.rows[0].cells[0].text = "What is your uptime SLA?"
    t1.rows[0].cells[1].text = "What is your security protocol?"
    # Row 1: Placeholders / Empty
    t1.rows[1].cells[0].text = "[Insert Response]"
    t1.rows[1].cells[1].text = ""  # Empty cell
    
    # 3. Table Horizontal Layout (Question cell to the left)
    doc.add_heading("Section 3: Security Details Table", level=2)
    t2 = doc.add_table(rows=2, cols=2)
    # Row 0
    t2.rows[0].cells[0].text = "What is your security protocol?"
    t2.rows[0].cells[1].text = ""  # Empty cell to fill
    # Row 1
    t2.rows[1].cells[0].text = "What is your uptime SLA?"
    t2.rows[1].cells[1].text = "[Insert response here]"
    
    doc.save(path)
    print(f"Mock template created at {path}")

def run_tests():
    load_dotenv()
    print("=== AutoRFP AI Batch Features Test ===")
    
    # 1. Test template filling logic locally (Offline)
    mock_template_path = "test_template.docx"
    create_mock_docx_template(mock_template_path)
    
    mock_qa_pairs = [
        {
            "question": "What is your security protocol?",
            "answer": "Our security protocol is SOC 2 Type II certified and complies with GDPR and CCPA regulations."
        },
        {
            "question": "What is your uptime SLA?",
            "answer": "We offer a 99.9% uptime SLA for all enterprise subscription plans, backed by redundant cloud hosting."
        }
    ]
    
    print("\nTesting template fill using fill_rfp_docx_template...")
    try:
        filled_bytes = fill_rfp_docx_template(mock_template_path, mock_qa_pairs)
        filled_doc = docx.Document(io.BytesIO(filled_bytes))
        
        # Verify paragraph replacement
        paragraphs_text = [p.text.strip() for p in filled_doc.paragraphs if p.text.strip()]
        print("\nVerifying paragraph replacements:")
        print("  Paragraphs found in filled doc:")
        for pt in paragraphs_text:
            print(f"    - {pt}")
            
        assert "SOC 2 Type II certified" in paragraphs_text[2], "Paragraph placeholder 1 not filled correctly!"
        assert "99.9% uptime SLA" in paragraphs_text[4], "Paragraph placeholder 2 not filled correctly!"
        print("  SUCCESS: Paragraph placeholders replaced correctly.")
        
        # Verify vertical table replacement
        print("\nVerifying vertical table layout:")
        t1 = filled_doc.tables[0]
        # Row 1 col 0 should be uptime SLA
        # Row 1 col 1 should be security protocol
        val1 = t1.rows[1].cells[0].text.strip()
        val2 = t1.rows[1].cells[1].text.strip()
        print(f"  t1 (vertical) Cell [1,0] text: {repr(val1)}")
        print(f"  t1 (vertical) Cell [1,1] text: {repr(val2)}")
        assert "99.9% uptime SLA" in val1, "Vertical table cell [1,0] not filled correctly!"
        assert "SOC 2 Type II certified" in val2, "Vertical table cell [1,1] not filled correctly!"
        print("  SUCCESS: Vertical table layout filled correctly.")
        
        # Verify horizontal table replacement
        print("\nVerifying horizontal table layout:")
        t2 = filled_doc.tables[1]
        # Row 0 col 1 should be security protocol
        # Row 1 col 1 should be uptime SLA
        val3 = t2.rows[0].cells[1].text.strip()
        val4 = t2.rows[1].cells[1].text.strip()
        print(f"  t2 (horizontal) Cell [0,1] text: {repr(val3)}")
        print(f"  t2 (horizontal) Cell [1,1] text: {repr(val4)}")
        assert "SOC 2 Type II certified" in val3, "Horizontal table cell [0,1] not filled correctly!"
        assert "99.9% uptime SLA" in val4, "Horizontal table cell [1,1] not filled correctly!"
        print("  SUCCESS: Horizontal table layout filled correctly.")
        
        print("\nOffline Template Filling Test: ALL PASSED!")
        
    except Exception as ex:
        print(f"  FAIL: Template filling test failed: {str(ex)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    finally:
        if os.path.exists(mock_template_path):
            os.remove(mock_template_path)
            
    # 2. Test Batch Ingestion & Extraction (Online)
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("\n[Online Tests] SKIPPED: GEMINI_API_KEY is not defined in .env")
        print("\nAll offline tests passed successfully!")
        return
        
    print("\n[Online Tests] Gemini API connection detected. Initializing database batch test...")
    try:
        client = get_chroma_client(db_path="test_batch_chroma_db")
        collection = get_or_create_collection(client, api_key, collection_name="test_batch_kb")
        
        # Create two simple documents to ingest in a batch
        doc1_path = "test_doc1.docx"
        doc2_path = "test_doc2.docx"
        
        d1 = docx.Document()
        d1.add_paragraph("Our company name is AutoRFP Technologies.")
        d1.add_paragraph("We were founded in 2024 to simplify procurement processes.")
        d1.save(doc1_path)
        
        d2 = docx.Document()
        d2.add_paragraph("We deploy all application instances using Kubernetes cluster deployments.")
        d2.add_paragraph("Our hosting provider is Google Cloud Platform (GCP) located in regional multi-zones.")
        d2.save(doc2_path)
        
        print(f"  Ingesting {doc1_path} and {doc2_path} in a single batch...")
        chunks_added = ingest_documents_batch([doc1_path, doc2_path], collection)
        print(f"  SUCCESS: Batch ingested {chunks_added} chunks into vector database.")
        
        # Run a batch query simulation
        questions = [
            "What is the name of your company?",
            "Where do you host your application?"
        ]
        print(f"  Running batch RAG pipeline for: {questions}")
        qa_results = batch_process_rfp_questions(questions, collection, api_key)
        
        print("\n--- BATCH PROCESSED RESULTS ---")
        for idx, item in enumerate(qa_results):
            print(f"Q{idx+1}: {item['question']}")
            print(f"A{idx+1}: {item['answer']}")
            print(f"Sources: {[s['source'] for s in item.get('sources', [])]}")
            print("-" * 30)
            
        assert len(qa_results) == 2, "Batch processing results count should be 2!"
        assert "AutoRFP Technologies" in qa_results[0]["answer"], "Company name response did not retrieve correct context!"
        assert "Google Cloud" in qa_results[1]["answer"] or "GCP" in qa_results[1]["answer"], "Hosting response did not retrieve correct context!"
        print("  SUCCESS: Batch RAG pipeline processed questions and retrieved context correctly.")
        
        # Clean up files & database
        import shutil
        os.remove(doc1_path)
        os.remove(doc2_path)
        print("\nCleaning up test database...")
        client.delete_collection("test_batch_kb")
        del client
        shutil.rmtree("test_batch_chroma_db", ignore_errors=True)
        print("  SUCCESS: Batch integration test completed successfully!")
        
    except Exception as e:
        print(f"  FAIL: Online test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        # Clean up files in case of failure
        for f in ["test_doc1.docx", "test_doc2.docx"]:
            if os.path.exists(f):
                os.remove(f)
        sys.exit(1)

if __name__ == "__main__":
    run_tests()
