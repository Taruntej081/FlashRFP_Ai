import io
import docx
from exporter import fill_rfp_docx_template, generate_docx_stream, generate_batch_docx_stream

def test_word_injection():
    print("--- Testing Word Injection & Format-Lock Template Engine ---")
    
    # 1. Create a sample Word template with various placeholder patterns and tables
    doc = docx.Document()
    
    doc.add_heading("Section 1: Information Security & Data Governance", level=1)
    
    p1 = doc.add_paragraph("Question 1.1: Describe your data encryption standards at rest and in transit.")
    p2 = doc.add_paragraph("[Insert Answer Here]")
    
    doc.add_heading("Section 2: Technical Compliance Table", level=2)
    
    table = doc.add_table(rows=3, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Requirement"
    hdr_cells[1].text = "Bidder Technical Response"
    
    row1 = table.rows[1].cells
    row1[0].text = "Does the solution support Single Sign-On (SSO) via SAML 2.0 / OAuth?"
    row1[1].text = "[Vendor Response]"
    
    row2 = table.rows[2].cells
    row2[0].text = "Provide SLA commitment for system uptime availability."
    row2[1].text = "" # Empty cell placeholder
    
    template_stream = io.BytesIO()
    doc.save(template_stream)
    template_bytes = template_stream.getvalue()
    print(f"Created sample Word template ({len(template_bytes)} bytes)")
    
    # 2. Mock QA pairs retrieved from RAG pipeline
    qa_pairs = [
        {
            "question": "Question 1.1: Describe your data encryption standards at rest and in transit.",
            "answer": "All data at rest is encrypted using FIPS 140-2 validated AES-256 bit encryption. Data in transit is secured via TLS 1.3 with automated certificate rotation."
        },
        {
            "question": "Does the solution support Single Sign-On (SSO) via SAML 2.0 / OAuth?",
            "answer": "Yes, FlashRFP natively supports Okta, Azure AD, PingIdentity, and custom SAML 2.0 / OAuth 2.0 identity providers with SCIM provisioning."
        },
        {
            "question": "Provide SLA commitment for system uptime availability.",
            "answer": "FlashRFP guarantees 99.95% system uptime SLA backed by 24/7 Enterprise Support and automated multi-region failover."
        }
    ]
    
    # 3. Test fill_rfp_docx_template
    filled_bytes = fill_rfp_docx_template(io.BytesIO(template_bytes), qa_pairs)
    assert len(filled_bytes) > 0, "Filled Word document stream is empty!"
    
    filled_doc = docx.Document(io.BytesIO(filled_bytes))
    
    full_text = "\n".join([p.text for p in filled_doc.paragraphs])
    table_text = "\n".join([cell.text for row in filled_doc.tables[0].rows for cell in row.cells])
    
    print("\nFilled Paragraphs Preview:")
    print(" ", full_text)
    print("\nFilled Table Preview:")
    print(" ", table_text)
    
    assert "AES-256 bit encryption" in full_text, "Paragraph placeholder not filled correctly!"
    assert "Okta, Azure AD" in table_text, "Table placeholder [Vendor Response] not filled!"
    assert "99.95% system uptime" in table_text, "Empty table cell not filled!"
    
    # 4. Test standalone DOCX generation
    docx_stream = generate_docx_stream("What is your backup policy?", "Automated daily snapshots with 30-day retention.", [])
    assert len(docx_stream) > 0, "Standalone DOCX stream generation failed!"
    
    print("\n[PASS] Word Injection & Format-Lock Template Engine PASSED cleanly!")

if __name__ == "__main__":
    test_word_injection()
