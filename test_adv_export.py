from exporter import generate_advanced_proposal_docx
import docx

def test_advanced_docx_generator():
    print("--- Testing Advanced Proposal Word Document Generator ---")
    
    extracted_data = [
        {
            "question": "1.1 Describe your architecture for N+1 redundancy.",
            "answer": "Our architecture guarantees N+1 redundancy across all critical paths...",
            "source": "Past_Proposal_SBI_2023.pdf (Page 14)",
            "manual_review": False
        },
        {
            "question": "1.2 Provide the PBG Bank Guarantee format.",
            "answer": "MANUAL REVIEW REQUIRED: No historical data found in the knowledge base.",
            "source": "No source found.",
            "manual_review": True
        }
    ]

    buffer = generate_advanced_proposal_docx(extracted_data, project_name="State Bank RFP Proposal")
    assert buffer is not None, "Buffer returned is None!"
    
    doc = docx.Document(buffer)
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Headings/Sections: {len(doc.headings if hasattr(doc, 'headings') else [])}")
    
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "State Bank RFP Proposal" in full_text
    assert "Technical & Compliance Proposal Draft" in full_text
    assert "1.1 Describe your architecture for N+1 redundancy." in full_text
    assert "MANUAL REVIEW REQUIRED" in full_text
    assert "Source Reference: Past_Proposal_SBI_2023.pdf (Page 14)" in full_text
    
    print("[PASS] Advanced Word Proposal Generator Test PASSED cleanly!")

if __name__ == "__main__":
    test_advanced_docx_generator()
