import os
import sys
import chromadb
import google.generativeai as genai
from dotenv import load_dotenv
from rag_engine import get_chroma_client, get_or_create_collection, ingest_document
import docx

def create_dummy_docx():
    doc = docx.Document()
    doc.add_heading("Standard Delivery Methodology", level=1)
    doc.add_paragraph(
        "Our standard delivery methodology utilizes an Agile framework divided into 2-week sprints. "
        "Each sprint begins with planning and ends with a review and retrospective."
    )
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Timeline'
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Discovery'
    row_cells[1].text = 'Weeks 1-2'
    doc.save("test_delivery.docx")

def run_test():
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("FAIL: No GEMINI_API_KEY found in .env")
        return
        
    create_dummy_docx()
    print("Dummy DOCX file created.")
    
    try:
        client = get_chroma_client()
        collection = get_or_create_collection(client, api_key)
        
        print("Ingesting DOCX file into local ChromaDB...")
        chunks = ingest_document("test_delivery.docx", collection)
        print(f"SUCCESS: Ingested {chunks} chunks!")
    except Exception as e:
        print(f"FAIL: Ingestion failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    finally:
        if os.path.exists("test_delivery.docx"):
            os.remove("test_delivery.docx")

if __name__ == "__main__":
    run_test()
